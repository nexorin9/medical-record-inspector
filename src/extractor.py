"""
文本提取模块 - Medical Record Inspector
从多种格式（PDF、Word、TXT）提取病历文本，支持中文文本清洗和标准化
"""

import re
import os
from pathlib import Path
from typing import Optional
from PyPDF2 import PdfReader
from docx import Document
import logging

logger = logging.getLogger(__name__)


def extract_text_from_pdf(pdf_path: str) -> str:
    """从 PDF 文件中提取文本

    Args:
        pdf_path: PDF 文件路径

    Returns:
        提取的文本内容
    """
    try:
        reader = PdfReader(pdf_path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        logger.info(f"成功从 PDF 提取 {len(text)} 字符")
        return text
    except Exception as e:
        logger.error(f"PDF 提取失败: {e}")
        raise


def extract_text_from_docx(docx_path: str) -> str:
    """从 Word (.docx) 文件中提取文本

    Args:
        docx_path: Word 文件路径

    Returns:
        提取的文本内容
    """
    try:
        doc = Document(docx_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        # 提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
        logger.info(f"成功从 Word 提取 {len(text)} 字符")
        return text
    except Exception as e:
        logger.error(f"Word 提取失败: {e}")
        raise


def extract_text_from_txt(txt_path: str) -> str:
    """从纯文本文件中提取文本

    Args:
        txt_path: 文本文件路径

    Returns:
        文件内容
    """
    try:
        with open(txt_path, 'r', encoding='utf-8') as f:
            text = f.read()
        logger.info(f"成功从 TXT 读取 {len(text)} 字符")
        return text
    except UnicodeDecodeError:
        # 尝试其他编码
        with open(txt_path, 'r', encoding='gbk') as f:
            text = f.read()
        logger.info(f"成功从 TXT 读取 (GBK 编码) {len(text)} 字符")
        return text
    except Exception as e:
        logger.error(f"TXT 读取失败: {e}")
        raise


def clean_text(text: str) -> str:
    """清洗文本 - 去除多余空白和特殊字符

    Args:
        text: 原始文本

    Returns:
        清洗后的文本
    """
    # 去除连续的空行（转换为单个空行）
    text = re.sub(r'\n\s*\n', '\n\n', text)
    # 去除连续多个空行（转换为最多两个连续换行）
    text = re.sub(r'\n{3,}', '\n\n', text)
    # 去除行首行尾的多余空白
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    # 去除连续多个空格
    text = re.sub(r' +', ' ', text)
    # 去除不能打印的字符
    text = ''.join(char for char in text if char.isprintable() or char in '\n\t')
    logger.debug(f"文本清洗: {len(text)} 字符")
    return text


def standardize_chinese(text: str) -> str:
    """中文标准化 - 全角转半角、繁体转简体

    Args:
        text: 原始文本

    Returns:
        标准化后的文本
    """
    # 全角转半角
    result = []
    for char in text:
        code = ord(char)
        # 全角空格
        if code == 12288:
            code = 32
        # 全角字符 (33-126)
        elif 65281 <= code <= 65374:
            code -= 65248
        result.append(chr(code) if code != 12288 else ' ')

    text = ''.join(result)

    # 注意：繁体转简体需要额外的库（如 opencc）
    # 这里只做简单的常见繁体字转换
    traditional_to_simplified = {
        '(','（',
        ')','）',
    }
    # 简化版本：只转换常见的全角标点
    text = text.replace('（', '(').replace('）', ')')
    text = text.replace('【', '[').replace('】', ']')
    text = text.replace('：', ':').replace('；', ';')

    logger.debug(f"中文标准化完成")
    return text


def extract_text(file_path: str) -> str:
    """自动检测文件类型并提取文本

    Args:
        file_path: 文件路径

    Returns:
        提取的文本内容
    """
    ext = Path(file_path).suffix.lower()

    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    elif ext == '.txt':
        return extract_text_from_txt(file_path)
    else:
        # 尝试所有方法
        logger.warning(f"未知文件类型 {ext}，尝试所有提取方法")
        try:
            return extract_text_from_txt(file_path)
        except:
            try:
                return extract_text_from_pdf(file_path)
            except:
                return extract_text_from_docx(file_path)


def extract_text_from_directory(dir_path: str) -> dict:
    """批量提取目录下所有文件的文本

    Args:
        dir_path: 目录路径

    Returns:
        字典 {文件名: 文本内容}
    """
    results = {}
    supported_extensions = ['.pdf', '.docx', '.txt']

    for filename in os.listdir(dir_path):
        ext = Path(filename).suffix.lower()
        if ext in supported_extensions:
            file_path = os.path.join(dir_path, filename)
            try:
                text = extract_text(file_path)
                results[filename] = text
                logger.info(f"成功处理文件: {filename}")
            except Exception as e:
                logger.error(f"处理文件失败 {filename}: {e}")

    return results


def save_text_to_file(text: str, output_path: str) -> None:
    """保存文本到文件

    Args:
        text: 文本内容
        output_path: 输出文件路径
    """
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)
    logger.info(f"文本保存到: {output_path}")


if __name__ == '__main__':
    # 简单测试
    logging.basicConfig(level=logging.INFO)

    # 创建测试文件
    test_dir = Path(__file__).parent.parent / 'data' / 'samples'
    test_dir.mkdir(parents=True, exist_ok=True)

    # 创建示例文本文件
    test_txt = test_dir / 'test_example.txt'
    test_txt.write_text("这是一个测试病历。\n主诉：发热咳嗽3天。\n诊断：肺炎。", encoding='utf-8')

    # 测试提取
    text = extract_text(str(test_txt))
    print(f"Extracted text length: {len(text)}")
    print(f"First 100 chars: {text[:100]}")

    # 测试清洗
    cleaned = clean_text(text)
    print(f"Cleaned text length: {len(cleaned)}")

    # 测试标准化
    standardized = standardize_chinese(text)
    print(f"Standardized text length: {len(standardized)}")
