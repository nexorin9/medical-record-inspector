"""
文件解析工具 - File Parser Utils
支持 DOCX 和 PDF 文件的文本提取
"""

import os
import tempfile
from typing import Tuple, Optional
from io import BytesIO

# DOCX 解析
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# PDF 解析
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


class FileParserError(Exception):
    """文件解析异常"""
    pass


class UnsupportedFileTypeError(FileParserError):
    """不支持的文件类型异常"""
    pass


class FileParser:
    """
    文件解析器

    支持格式：
    - .docx: Microsoft Word 文档
    - .pdf: PDF 文档
    """

    SUPPORTED_EXTENSIONS = {'.docx', '.pdf'}
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

    @staticmethod
    def parse_file(file_path: str) -> Tuple[str, str]:
        """
        解析文件内容

        Args:
            file_path: 文件路径

        Returns:
            Tuple[str, str]: (文件扩展名, 提取的文本内容)

        Raises:
            UnsupportedFileTypeError: 不支持的文件类型
            FileParserError: 解析失败
        """
        ext = os.path.splitext(file_path)[1].lower()

        if ext not in FileParser.SUPPORTED_EXTENSIONS:
            raise UnsupportedFileTypeError(
                f"不支持的文件类型: {ext}. "
                f"支持的格式: {', '.join(FileParser.SUPPORTED_EXTENSIONS)}"
            )

        if ext == '.docx':
            return ext, FileParser._parse_docx(file_path)
        elif ext == '.pdf':
            return ext, FileParser._parse_pdf(file_path)

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        """
        解析 DOCX 文件

        Args:
            file_path: DOCX 文件路径

        Returns:
            提取的文本内容

        Raises:
            FileParserError: 解析失败
        """
        if not DOCX_AVAILABLE:
            raise FileParserError(
                "python-docx 未安装。请运行: pip install python-docx"
            )

        try:
            doc = Document(file_path)
            paragraphs = []

            # 提取段落文本
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())

            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(' | '.join(row_text))

            return '\n\n'.join(paragraphs)

        except Exception as e:
            raise FileParserError(f"DOCX 解析失败: {str(e)}")

    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        """
        解析 PDF 文件

        Args:
            file_path: PDF 文件路径

        Returns:
            提取的文本内容

        Raises:
            FileParserError: 解析失败
        """
        if not PDF_AVAILABLE:
            raise FileParserError(
                "PyPDF2 未安装。请运行: pip install PyPDF2"
            )

        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                pages = []

                for page in reader.pages:
                    text = page.extract_text()
                    if text and text.strip():
                        pages.append(text.strip())

                return '\n\n'.join(pages)

        except Exception as e:
            raise FileParserError(f"PDF 解析失败: {str(e)}")

    @staticmethod
    def parse_file_binary(file_content: bytes, filename: str) -> Tuple[str, str]:
        """
        从二进制内容解析文件

        Args:
            file_content: 文件二进制内容
            filename: 文件名

        Returns:
            Tuple[str, str]: (文件扩展名, 提取的文本内容)
        """
        ext = os.path.splitext(filename)[1].lower()

        if ext not in FileParser.SUPPORTED_EXTENSIONS:
            raise UnsupportedFileTypeError(
                f"不支持的文件类型: {ext}. "
                f"支持的格式: {', '.join(FileParser.SUPPORTED_EXTENSIONS)}"
            )

        # 创建临时文件
        with tempfile.NamedTemporaryFile(
            suffix=ext, delete=False, mode='wb'
        ) as tmp_file:
            tmp_file.write(file_content)
            tmp_path = tmp_file.name

        try:
            return FileParser.parse_file(tmp_path)
        finally:
            # 清理临时文件
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)

    @staticmethod
    def validate_file(file_name: str, file_size: Optional[int] = None) -> None:
        """
        验证文件

        Args:
            file_name: 文件名
            file_size: 文件大小（字节），可选

        Raises:
            UnsupportedFileTypeError: 不支持的文件类型
            FileParserError: 文件过大
        """
        ext = os.path.splitext(file_name)[1].lower()

        if ext not in FileParser.SUPPORTED_EXTENSIONS:
            raise UnsupportedFileTypeError(
                f"不支持的文件类型: {ext}. "
                f"支持的格式: {', '.join(FileParser.SUPPORTED_EXTENSIONS)}"
            )

        if file_size and file_size > FileParser.MAX_FILE_SIZE:
            raise FileParserError(
                f"文件过大。最大支持 10MB，当前: {file_size / 1024 / 1024:.2f}MB"
            )


def parse_file(file_path: str) -> str:
    """
    便捷函数：解析文件并返回文本内容

    Args:
        file_path: 文件路径

    Returns:
        提取的文本内容

    Raises:
        FileParserError: 解析失败
    """
    _, text = FileParser.parse_file(file_path)
    return text


def parse_file_binary(file_content: bytes, filename: str) -> str:
    """
    便捷函数：从二进制内容解析文件并返回文本

    Args:
        file_content: 文件二进制内容
        filename: 文件名

    Returns:
        提取的文本内容

    Raises:
        FileParserError: 解析失败
    """
    _, text = FileParser.parse_file_binary(file_content, filename)
    return text
